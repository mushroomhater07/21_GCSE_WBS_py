import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define a function to select the input file
def select_input_file():
    input_path_var.set("%userprofile%/Documents/BiRP NEA/Assets/Scripts")
    file_path = filedialog.askdirectory(initialdir=r"C:/Users/hey/Documents/BiRP NEA/Assets/Scripts")
    # filetypes=[('C# Files', '*.cs'),('other','*.*')]
    input_path_var.set(file_path)

# Define a function to select the output file
def select_output_file():
    output_path_var.set("%userprofile%/Desktop/Script.docx")
    file_path = filedialog.asksaveasfilename(initialdir=r"C:/Users/hey/Desktop",initialfile=r"Script.docx",filetypes=[('Word Documents', '*.docx')])
    output_path_var.set(file_path)

# Set up the GUI
root = tk.Tk()
root.title('Convert C# Source Code to Word Document')

# Set up the input file selection frame
input_frame = tk.Frame(root)
input_frame.pack(fill=tk.X, padx=10, pady=5)

input_label = tk.Label(input_frame, text='Input File:')
input_label.pack(side=tk.LEFT)

input_path_var = tk.StringVar()
input_path_entry = tk.Entry(input_frame, textvariable=input_path_var, width=50)
input_path_entry.pack(side=tk.LEFT, padx=5)

input_button = tk.Button(input_frame, text='Browse', command=select_input_file)
input_button.pack(side=tk.LEFT)

# Set up the output file selection frame
output_frame = tk.Frame(root)
output_frame.pack(fill=tk.X, padx=10, pady=5)

output_label = tk.Label(output_frame, text='Output File:')
output_label.pack(side=tk.LEFT)

output_path_var = tk.StringVar()
output_path_entry = tk.Entry(output_frame, textvariable=output_path_var, width=50)
output_path_entry.pack(side=tk.LEFT, padx=5)

output_button = tk.Button(output_frame, text='Browse', command=select_output_file)
output_button.pack(side=tk.LEFT)

# Define a function to convert the C# source code to a Word document
def convert_to_word_document():
    # Get the input and output file paths
    input_path = input_path_var.get()
    output_path = output_path_var.get()
    
    # Create a new Word document
    document = Document()
    
    # Recursively traverse the directory tree
    for root, dirs, files in os.walk(input_path):
        for file_name in files:
            if file_name.endswith('.cs'):
                # Read the contents of the file
                with open(os.path.join(root, file_name), 'r') as file:
                    content = file.read()
                
                # Add the file name as a heading
                # heading = document.add_heading(os.path.abspath(file_name), level=1)
                heading = document.add_heading(os.path.basename(file_name), level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add the contents of the file to a new paragraph in the document
                paragraph = document.add_paragraph()
                paragraph.add_run(content)
            
    # Save the document to the output path
    document.save(output_path)
    print(f'Successfully created document at {output_path}')

# Set up the conversion button
convert_button = tk.Button(root, text='Convert', command=convert_to_word_document)
convert_button.pack(padx=10, pady=10)

# Run the GUI
root.mainloop()
